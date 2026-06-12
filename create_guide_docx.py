from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('GridX: Watt & Byte - 베타 테스터 가이드', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('환영합니다! GridX: Watt & Byte는 전 세계 주요 데이터센터(DC)와 핵심 전력망(발전소) 데이터를 한눈에 조망하고, AI 기반의 인사이트를 제공하기 위해 개발된 통합 인프라 대시보드입니다.\n\n베타 테스트 기간 동안 아래의 핵심 기능들을 자유롭게 탐색해 보시고 많은 피드백 부탁드립니다.')

doc.add_heading('1. 글로벌 인프라 공간 시각화 (Interactive 3D-like Map)', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('데이터센터 & 발전소 맵핑: ').bold = True
p.add_run('전 세계에 흩어진 하이퍼스케일 데이터센터와 50MW 이상의 주요 발전소들의 위치를 인터랙티브 지도 위에 버블 형태로 직관적으로 보여줍니다.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('스마트 줌 & 클러스터링: ').bold = True
p.add_run('마우스 휠이나 확대/축소 버튼을 통해 지도를 탐색할 수 있습니다. 줌인 시 겹쳐있는 인프라가 자동으로 적당한 거리를 두고 분산되어 개별 시설을 정확히 클릭하고 확인할 수 있습니다.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('반응형 모바일 지원: ').bold = True
p.add_run('PC뿐만 아니라 모바일 브라우저에서도 화면 깨짐 없이 완벽하게 동작하도록 최적화되어 있습니다.')

doc.add_heading('2. 직관적인 필터링 시스템', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('인프라 토글 버튼: ').bold = True
p.add_run('상단의 [데이터센터] / [발전소] 버튼을 클릭해 원하는 종류의 인프라만 지도에 표시할 수 있습니다.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('연료별 발전소 멀티 셀렉트: ').bold = True
p.add_run('발전소 뷰 모드에서는 화면 상단에 슬라이드 메뉴가 나타납니다. 원자력, 태양광, 풍력, 석탄 등 발전소 종류를 중복 선택(Multi-select)하여 특정 에너지만 필터링해 분석할 수 있습니다.')

doc.add_heading('3. 실시간 뉴스 스크랩 및 키워드 추출', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('글로벌 인프라 뉴스 전광판: ').bold = True
p.add_run('상단 Ticker를 통해 해외 최신 데이터센터 및 AI 인프라 영문 뉴스를 한국어로 실시간 자동 번역하여 제공합니다.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('뉴스 자동 아카이빙: ').bold = True
p.add_run('사용자가 접속할 때마다 최신 기사를 불러와 브라우저에 최대 100개까지 자동 누적 저장합니다.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('핵심 키워드 뱃지: ').bold = True
p.add_run('[General Information] -> [지난뉴스 보기]를 클릭하면 그동안 수집된 뉴스 목록과 함께 3~4개의 주요 핵심 키워드가 해시태그(#) 형태로 추출되어 기사의 맥락을 바로 파악할 수 있습니다.')

doc.add_heading('4. AI 챗봇 원클릭 요약 연동', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('지능형 어시스턴트 내장: ').bold = True
p.add_run('우측 하단(또는 모바일 하단)의 챗봇 패널을 통해 "미국의 데이터센터 개수는?", "가장 큰 발전소는 어디야?" 같은 자연어 질문을 던지고 즉각적인 답변을 받을 수 있습니다.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('뉴스 원클릭 요약: ').bold = True
p.add_run('뉴스 목록 우측에 있는 [챗봇 요약] 버튼을 클릭하면, 기사 링크가 챗봇으로 자동 전송되며 복잡한 기사의 핵심 내용을 AI가 즉시 요약해 줍니다.')

doc.add_heading('5. 인프라 지식 베이스 (General Information)', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('시장 최신 동향 제공: ').bold = True
p.add_run('글로벌 빅테크 데이터센터 주요 지표(PUE, WUE 등) 및 최신 데이터센터 냉각 기술 동향(공랭식, D2C 수랭식, 액침 냉각) 등 전문적인 백그라운드 지식을 제공합니다.')

doc.add_paragraph('\n[테스트 팁]')
p = doc.add_paragraph('지도를 최대한 확대해 시설 간 겹침 현상이 잘 해소되는지 확인해 보시고, 뉴스 목록에서 [지난뉴스 보기]와 [챗봇 요약] 기능이 매끄럽게 연동되는지 중점적으로 테스트해 주세요!')
p.runs[0].font.color.rgb = RGBColor(0x38, 0xbd, 0xf8)

doc.save('GridX_Beta_Tester_Guide.docx')
